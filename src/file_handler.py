"""文件处理模块"""

import csv
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from .i18n import i18n


@dataclass
class FileContent:
    """文件内容"""
    path: str
    name: str
    format: str
    content: str
    size: int
    error: Optional[str] = None


class FileHandler:
    """文件处理器"""

    MAX_FILE_SIZE = 500_000  # 500KB 字符限制

    SUPPORTED_FORMATS = {
        ".txt": "text",
        ".md": "markdown",
        ".csv": "csv",
        ".pdf": "pdf",
        ".docx": "docx",
        ".json": "json",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".py": "code",
        ".js": "code",
        ".ts": "code",
        ".java": "code",
        ".cpp": "code",
        ".c": "code",
        ".go": "code",
        ".rs": "code",
        ".rb": "code",
        ".html": "code",
        ".css": "code",
        ".xml": "code",
        ".sql": "code",
        ".sh": "code",
        ".bat": "code",
        ".log": "text",
    }

    def load_file(self, file_path: str) -> FileContent:
        """加载文件内容

        Args:
            file_path: 文件路径

        Returns:
            FileContent对象
        """
        path = Path(file_path).expanduser().resolve()

        if not path.exists():
            return FileContent(
                path=str(path),
                name=path.name,
                format="unknown",
                content="",
                size=0,
                error=i18n.t("file_not_found", path=str(path)),
            )

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            return FileContent(
                path=str(path),
                name=path.name,
                format=suffix,
                content="",
                size=0,
                error=i18n.t("file_unsupported", format=suffix),
            )

        format_type = self.SUPPORTED_FORMATS[suffix]
        try:
            if format_type == "pdf":
                content = self._read_pdf(path)
            elif format_type == "docx":
                content = self._read_docx(path)
            elif format_type == "csv":
                content = self._read_csv(path)
            elif format_type in ("yaml", "yml"):
                content = self._read_yaml(path)
            elif format_type == "json":
                content = self._read_json(path)
            else:
                content = self._read_text(path)
        except Exception as e:
            return FileContent(
                path=str(path),
                name=path.name,
                format=format_type,
                content="",
                size=0,
                error=i18n.t("file_error", error=str(e)),
            )

        size = len(content)
        if size > self.MAX_FILE_SIZE:
            return FileContent(
                path=str(path),
                name=path.name,
                format=format_type,
                content="",
                size=size,
                error=i18n.t("file_too_large", size=size, limit=self.MAX_FILE_SIZE),
            )

        return FileContent(
            path=str(path),
            name=path.name,
            format=format_type,
            content=content,
            size=size,
        )

    def _read_text(self, path: Path) -> str:
        """读取文本文件"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError(f"无法解码文件: {path}")

    def _read_csv(self, path: Path) -> str:
        """读取CSV文件并转为可读文本"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        content = None
        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if content is None:
            raise ValueError(f"无法解码CSV文件: {path}")

        reader = csv.reader(io.StringIO(content))
        rows = list(reader)
        if not rows:
            return ""

        # 格式化为可读表格
        lines = []
        header = rows[0]
        lines.append(" | ".join(header))
        lines.append("-" * len(lines[0]))
        for row in rows[1:]:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    def _read_pdf(self, path: Path) -> str:
        """读取PDF文件"""
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(f"--- Page {i + 1} ---\n{text}")
                return "\n\n".join(pages)
        except ImportError:
            raise ValueError("PDF support requires PyPDF2: pip install PyPDF2")

    def _read_docx(self, path: Path) -> str:
        """读取DOCX文件"""
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # 也读取表格
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except ImportError:
            raise ValueError("DOCX support requires python-docx: pip install python-docx")

    def _read_json(self, path: Path) -> str:
        """读取JSON文件"""
        import json
        text = self._read_text(path)
        try:
            data = json.loads(text)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            return text

    def _read_yaml(self, path: Path) -> str:
        """读取YAML文件"""
        import yaml
        text = self._read_text(path)
        try:
            data = yaml.safe_load(text)
            return yaml.dump(data, default_flow_style=False, allow_unicode=True)
        except yaml.YAMLError:
            return text

    def get_context_prompt(self, file_content: FileContent) -> str:
        """生成可注入对话上下文的文件内容提示"""
        if file_content.error:
            return f"[File Error: {file_content.error}]"

        header = i18n.t("file_content_header", name=file_content.name)
        return f"{header}{file_content.content}"

    def list_supported_formats(self) -> list[str]:
        """列出支持的文件格式"""
        return sorted(set(self.SUPPORTED_FORMATS.values()))

    def is_supported(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        suffix = Path(file_path).suffix.lower()
        return suffix in self.SUPPORTED_FORMATS


# 全局实例
file_handler = FileHandler()
